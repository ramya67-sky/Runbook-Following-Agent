import json
import csv
import io
import re
from html.parser import HTMLParser

# Try importing third-party libraries; handle ImportError if not installed
try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import docx
except ImportError:
    docx = None

try:
    import openpyxl
except ImportError:
    openpyxl = None

try:
    import yaml
except ImportError:
    yaml = None


class MLStripper(HTMLParser):
    def __init__(self):
        super().__init__()
        self.reset()
        self.strict = False
        self.convert_charrefs = True
        self.text = io.StringIO()

    def handle_data(self, d):
        self.text.write(d)

    def get_data(self):
        return self.text.getvalue()


def strip_tags(html: str) -> str:
    s = MLStripper()
    s.feed(html)
    return s.get_data()


def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """
    Routes to the correct parser based on filename extension.
    Supported: .md, .txt, .pdf, .docx, .xlsx, .json, .yaml, .yml, .csv, .rst, .html, .htm
    Returns plain text string.
    """
    if not filename:
        return file_bytes.decode('utf-8', errors='replace')

    ext = filename.split('.')[-1].lower()

    if ext in ('txt', 'md', 'rst'):
        try:
            return file_bytes.decode('utf-8')
        except UnicodeDecodeError:
            try:
                return file_bytes.decode('latin-1')
            except Exception:
                return file_bytes.decode('utf-8', errors='replace')

    elif ext == 'pdf':
        if not pypdf:
            raise ImportError("pypdf is not installed on the system.")
        pdf_file = io.BytesIO(file_bytes)
        reader = pypdf.PdfReader(pdf_file)
        text_parts = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text_parts.append(t)
        return "\n".join(text_parts)

    elif ext == 'docx':
        if not docx:
            raise ImportError("python-docx is not installed on the system.")
        doc_file = io.BytesIO(file_bytes)
        doc = docx.Document(doc_file)
        text_parts = []
        for paragraph in doc.paragraphs:
            text_parts.append(paragraph.text)
        # Parse tables as well
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                text_parts.append(" | ".join(row_text))
        return "\n".join(text_parts)

    elif ext == 'xlsx':
        if not openpyxl:
            raise ImportError("openpyxl is not installed on the system.")
        excel_file = io.BytesIO(file_bytes)
        wb = openpyxl.load_workbook(excel_file, read_only=True, data_only=True)
        text_parts = []
        for sheet in wb.worksheets:
            text_parts.append(f"--- Sheet: {sheet.title} ---")
            for row in sheet.iter_rows(values_only=True):
                row_str = " | ".join(str(val) if val is not None else "" for val in row)
                if row_str.replace("|", "").strip():
                    text_parts.append(row_str)
        return "\n".join(text_parts)

    elif ext == 'json':
        content = file_bytes.decode('utf-8', errors='replace')
        try:
            parsed = json.loads(content)
            return json.dumps(parsed, indent=2)
        except Exception:
            return content

    elif ext in ('yaml', 'yml'):
        if not yaml:
            raise ImportError("pyyaml is not installed on the system.")
        content = file_bytes.decode('utf-8', errors='replace')
        try:
            parsed = yaml.safe_load(content)
            return yaml.dump(parsed, default_flow_style=False)
        except Exception:
            return content

    elif ext == 'csv':
        content = file_bytes.decode('utf-8', errors='replace')
        f = io.StringIO(content)
        reader = csv.reader(f)
        text_parts = []
        for row in reader:
            text_parts.append(" | ".join(row))
        return "\n".join(text_parts)

    elif ext in ('html', 'htm'):
        content = file_bytes.decode('utf-8', errors='replace')
        try:
            return strip_tags(content)
        except Exception:
            return content

    else:
        # Unknown extension, fallback to text decode
        try:
            return file_bytes.decode('utf-8')
        except UnicodeDecodeError:
            try:
                return file_bytes.decode('latin-1')
            except Exception:
                return file_bytes.decode('utf-8', errors='replace')
